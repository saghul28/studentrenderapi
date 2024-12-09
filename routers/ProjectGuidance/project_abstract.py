from fastapi import APIRouter,Form
from pydantic import BaseModel
from docx import Document
from fastapi.responses import FileResponse
import google.generativeai as genai
import os
from dotenv import load_dotenv

load_dotenv()

def mygemai(prompt):
    client =  genai.configure(api_key=os.getenv("GENAIAPI_KEY"))
    model =   genai.GenerativeModel('gemini-1.5-flash')
    chat = model.start_chat()
    response = chat.send_message(prompt)
    final_response = response.text.replace("#","").replace("*","")
    return final_response
    


router = APIRouter() 



@router.post('/')
def generatingAbstrct(projectName:str=Form(...),degree:str=Form(...),department:str=Form(...),projectDescription:str=Form(...),frontend:str=Form(...),backend:str=Form(...)):
   
    prompt = (
    f"Create a comprehensive and detailed project abstract for a {degree} degree in {department}. "
    f"The project is titled '{projectName}', and its goal is to {projectDescription}. The abstract should be "
    f"divided into clearly defined sections, with a minimum of three pages of content. Each section should be "
    f"thorough and elaborate to ensure a deep understanding of the project.\n\n"

    f"1. **Introduction**:\n"
    f"   - Start with an introduction to the project, including the problem being solved and its relevance to the field. "
    f"   - Provide background information about the problem, its significance, and the motivation behind this project. "
    f"   - Mention why this project is important, its potential impact, and its scope.\n\n"

    f"2. **Project Details**:\n"
    f"   - Describe the overall objectives and goals of the project. "
    f"   - Provide a detailed breakdown of the methodologies used and the expected outcomes. "
    f"   - Explain how this project is innovative or unique compared to similar projects. "
    f"   - Include the key technologies involved in the project and how they contribute to achieving the project goals.\n\n"

    f"3. **Frontend**:\n"
    f"   - Describe the frontend technologies used, such as frameworks and libraries. "
    f"   - Explain why these technologies were chosen for the project. "
    f"   - Discuss how the frontend provides a user-friendly experience and supports the project's objectives. "
    f"   - Highlight key features like responsiveness, interactivity, or accessibility.\n\n"

    f"4. **Backend**:\n"
    f"   - Detail the backend technologies used, including frameworks, databases, and server-side technologies. "
    f"   - Discuss the reasons for choosing these technologies, focusing on scalability, performance, and security. "
    f"   - Explain how the backend handles data processing, storage, and integration with the frontend.\n\n"

    f"5. **System Recommendations**:\n"
    f"   - Provide suggestions for improving the system in the future, such as enhancing performance, adding new features, "
    f"     or improving security. "
    f"   - Recommend strategies for scaling the project as usage increases, ensuring sustainability and reliability. "
    f"   - Discuss potential areas for further research or development, including emerging technologies that could improve the system.\n\n"

    f"Ensure that the abstract is structured, technical, and written in a professional tone. "
    f"Make sure each section is detailed enough to fill at least one full page, focusing on providing a deep understanding "
    f"of the project’s goals, technologies, and future directions."
    )
    final_response = mygemai(prompt)
    doc = Document()
    doc.add_heading(f'Project Abstract:{projectName}',0)
    doc.add_paragraph(final_response.replace("#","").replace("*",""))
    
    file_path = f"project_abstract_{projectName}.docx"
    doc.save(file_path)
    return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=file_path)


@router.post("/generate_projects_steps")
def generateSteps(projectName:str=Form(...),degree:str=Form(...),department:str=Form(...),projectDescription:str=Form(...),frontend:str=Form(...),backend:str=Form(...)):
    prompt = (
    f"Create a detailed, step-by-step guide for a project titled '{projectName}', designed for a {degree} degree in {department}. "
    f"The goal of the project is to {projectDescription}. The guide should be divided into clearly defined sections, with each step providing a detailed explanation of the process, technologies used, code structure, and challenges faced. "
    f"The final document should be comprehensive, technical, and cover at least five pages.\n\n"

    f"1. **Introduction**:\n"
    f"   - Provide a brief overview of the project, its objectives, and the technologies that will be used. \n"
    f"   - Describe the project scope and explain its relevance to the field of {department}.\n"
    f"   - Define the importance of each phase of the project.\n\n"

    f"2. **Planning Stage**:\n"
    f"   - Detail the initial planning and research phase. Discuss how the requirements for the project were gathered.\n"
    f"   - Explain how the project objectives were established and refined.\n"
    f"   - Include information on the timeline, resources needed, and the initial feasibility study.\n"
    f"   - Software Declaration: Include the tools and platforms used for project management (e.g., Trello, Jira, etc.).\n\n"

    f"3. **Design Stage**:\n"
    f"   - Discuss the design phase in detail. Explain the system architecture, user interface, and user experience design.\n"
    f"   - Describe how the frontend and backend components were planned, including wireframes, database schemas, and system workflows.\n"
    f"   - Provide the **initial folder structure** for the project. For example, describe the directories for the frontend, backend, and database scripts.\n"
    f"     Example folder structure:\n"
    f"     - /frontend: Contains all frontend-related code (React components, stylesheets, etc.)\n"
    f"     - /backend: Contains all backend-related code (API routes, models, controllers, etc.)\n"
    f"     - /database: Contains scripts for database schema and migrations\n"
    f"   - Software Declaration: List the design tools used, such as Figma, Adobe XD, or similar tools.\n\n"

    f"4. **Development Stage**:\n"
    f"   - Break down the development phase into smaller steps, describing how each component of the system (frontend, backend, database) was developed.\n"
    f"   - **Database Structure**: Provide a clear description of the database schema, including tables, relationships, and any indexes used. For example, describe how users are stored in a 'users' table with fields for user ID, name, email, and password.\n"
    f"     Example database structure:\n"
    f"     - **Users table**: user_id (Primary Key), name, email, password\n"
    f"     - **Projects table**: project_id (Primary Key), user_id (Foreign Key), project_name, description\n"
    f"   - Provide examples of **initial code** that needs to be written. For example, creating a basic API with FastAPI, including a route for user registration, and integrating it with the database.\n"
    f"   - Provide an overview of the development process for both the frontend and backend, explaining how technologies like {frontend} (e.g., React, Vue.js, Angular) and {backend} (e.g., Flask, FastAPI, Django) were integrated.\n"
    f"   - Discuss key features such as authentication, data processing, and any major challenges encountered during the development phase.\n"
    f"   - Software Declaration: Clearly mention all programming languages, frameworks, libraries, and tools used in development. For example:\n"
    f"     - Frontend: {frontend}\n"
    f"     - Backend: {backend}\n"
    f"     - Database: MySQL/PostgreSQL\n"
    f"     - Other Tools: Git, Docker, etc.\n\n"

    f"5. **Testing and Quality Assurance**:\n"
    f"   - Describe the testing process in detail, including unit tests, integration tests, and user acceptance tests.\n"
    f"   - Discuss how the system was evaluated for scalability, performance, and security.\n"
    f"   - Mention any tools or libraries used for testing and debugging.\n"
    f"   - Software Declaration: Include tools like Postman for API testing, Selenium for UI testing, or Jest for unit testing.\n\n"

    f"6. **Deployment Stage**:\n"
    f"   - Explain the deployment process, including how the application was deployed to production.\n"
    f"   - Describe the steps taken for setting up servers, configuring cloud services, and ensuring continuous deployment.\n"
    f"   - Highlight any configuration management tools or scripts used for automation.\n"
    f"   - Software Declaration: Mention the deployment platforms or tools used, such as AWS, Heroku, Docker, or Kubernetes.\n\n"

    f"7. **Maintenance and Updates**:\n"
    f"   - Provide a plan for ongoing maintenance, bug fixes, and feature updates after the project has been deployed.\n"
    f"   - Explain how the system will be monitored for performance and errors, and how updates will be rolled out.\n"
    f"   - Discuss the role of version control (e.g., Git) in managing code updates.\n\n"

    f"8. **Conclusion**:\n"
    f"   - Summarize the project steps and how each phase contributed to achieving the project’s goals.\n"
    f"   - Discuss the lessons learned during the process and potential improvements for future iterations.\n"
    f"   - Provide recommendations for further development or research.\n\n"

    f"Ensure that each section is detailed enough to provide a full understanding of the project’s process, including the technologies, code structure, folder structure, and tools used at each stage. Focus on delivering a well-structured and technical document that can span at least five pages.\n\n"

    f"The result should be a step-by-step guide that not only outlines the phases of the project but also provides clarity on the technologies, code setup, and tools used at each stage of development."
    )

    final_response = mygemai(prompt)
    doc = Document()
    doc.add_heading(f'Project Steps:{projectName}',0)
    doc.add_paragraph(final_response.replace("#","").replace("*",""))
    
    file_path = f"project_Steps_{projectName}.docx"
    doc.save(file_path)
    return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=file_path)


@router.post("/generate_report")
def project_report(projectName:str=Form(...),degree:str=Form(...),department:str=Form(...),projectDescription:str=Form(...),frontend:str=Form(...),backend:str=Form(...)):
    prompt = (
    f"Generate a detailed report for the project titled '{projectName} and {projectDescription} "
    f"for small and medium-sized enterprises (SMEs). The report should have the following structure:\n\n"

    f"**Index:**\n\n"
    
    f"1. **Introduction**\n"
    f"   - Project Background\n"
    f"   - Problem Statement\n"
    f"   - Objectives\n"
    f"   - Technologies Overview ({frontend}, {backend}, PostgreSQL)\n\n"
    
    f"2. **Project Planning and Requirements**\n"
    f"   - Project Scope and Requirements\n"
    f"   - Technology Stack Selection (Why {frontend}, {backend})\n"
    f"   - User Stories/Use Cases\n\n"
    
    f"3. **System Design and Architecture**\n"
    f"   - High-Level System Architecture ({frontend}/{backend} interaction, API design)\n"
    f"   - Database Design (PostgreSQL schema for inventory tracking)\n"
    f"   - API Design ({backend} endpoints, RESTful routes)\n\n"
    
    f"4. **Frontend Development**\n"
    f"   - {frontend}\n"
    f"   - State Management about {frontend}\n"
    f"   - User Interface Design and Responsiveness\n"
    f"   - Code Snippets for Frontend Features \n\n"
    
    f"5. **Backend Development**\n"
    f"   - {backend} Overview ( {backend} details )\n"
    f"   - Database Interaction\n"
    f"   - API Design and Security (JWT Authentication, API Routes)\n"
    
    
    f"6. **Testing and Quality Assurance**\n"
    f"   - Testing Strategy (Unit, Integration Testing) in {frontend}{backend}\n"
    f"   - Tools Used \n"
    f"   - Code Coverage and Bug Fixes\n\n"
    
    f"7. **Deployment**\n"
    f"   - Deployment Platform (Heroku, AWS EC2, or Docker)\n"
    f"   - CI/CD Pipeline (GitHub Actions, Docker)\n"
    f"   - Steps for Deployment\n\n"
    
    f"8. **Challenges and Solutions**\n"
    f"   - Key Challenges (e.g., managing database consistency, API performance)\n"
    f"   - Solutions Implemented (e.g., database indexing, caching)\n\n"
    
    f"9. **Conclusion**\n"
    f"   - Project Summary\n"
    f"   - Lessons Learned\n"
    f"   - Future Improvements (e.g., adding more features, scaling for larger enterprises)\n\n"
    
    f"Please generate the content for each section in detail, with each subtopic having 300-400 words. "
    f"Ensure that the first page contains only the index with the listed subtopics. "
    f"After the index, generate detailed content for each subtopic, ensuring the report is informative and technical, "
    f"covering all aspects of the project development cycle, including system design, {backend}, {frontend}, testing, "
    f"deployment, and challenges faced."
    )
    final_response = mygemai(prompt)
    doc = Document()
    doc.add_heading(f'Project Abstract:{projectName}',0)
    doc.add_paragraph(final_response.replace("#","").replace("*",""))
    
    file_path = f"project_abstract_{projectName}.docx"
    doc.save(file_path)
    return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=file_path)
