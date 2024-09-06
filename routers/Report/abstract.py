from fastapi import FastAPI, Form, HTTPException,APIRouter
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
import os
import google.generativeai as genai

# app = FastAPI()
router =APIRouter()

# Configure Google GenAI (Remember to set your API key as an environment variable)
genai.configure(api_key=os.getenv("GENAIAPI_KEY"))

# Choose the model
model = genai.GenerativeModel('gemini-1.5-flash')


class ProjectIdea(BaseModel):
    idea: str

@router.post("/generate_abstract/")
async def generate_abstract(idea: str = Form(...)):
    try:
        # Use OpenAI to generate the abstract
        prompt = f"Generate a 150-word abstract for the following project idea:\n\n{idea}"
        chat = model.start_chat(history=[])
        response = chat.send_message(prompt)
        abstract_text = response.text.strip()

        # Create a Word document with the abstract
        document = Document()
        document.add_heading('Project Abstract', level=1)
        document.add_paragraph(abstract_text)
        file_path = "abstract.docx"
        document.save(file_path)

        return {"abstract": abstract_text, "file_path": file_path}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/download_abstract/")
async def download_abstract():
    file_path = "abstract.docx"
    if os.path.exists(file_path):
        return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename="abstract.docx")
    else:
        raise HTTPException(status_code=404, detail="File not found")
