from fastapi import APIRouter, HTTPException,Form
from pydantic import BaseModel
from docx import Document
import google.generativeai as genai  # Import google.generativeai
import io
from fastapi.responses import FileResponse, StreamingResponse
import  os


router = APIRouter()

genai_api_key = os.getenv("GENAIAPI_KEY")
genai.configure(api_key=genai_api_key)


# class ReportRequest(BaseModel):
#     project_title: str
#     project_description: str
#     project_type: str
#     pages: int
#     font: str


# @router.post("/generate-report/")
# async def generate_report(request: ReportRequest):
#     # Generate a specific prompt based on the project type and number of pages
#     prompt = (
#         f"Generate a detailed report for a {request.project_type} project titled '{request.project_title}'. "
#         f"The report should include the following description: '{request.project_description}'. "
#         f"The report should be around {request.pages} pages long."
#     )


#     try:
#         response = genai.generate_text(
#             prompt=prompt,
#             max_output_tokens=request.pages * 150  # Approx. 150 words per page
#         )
#         print(response)
#         generated_text = response.candidates[0]['output']
#     except Exception as e:
#         raise HTTPException(status_code=500, detail="Error with Generative AI API: " + str(e))


#     doc = Document()
#     doc.add_heading(request.project_title, 0)  # Add project title
#     doc.add_paragraph(f"Project Description: {request.project_description}")  # Add project description
#     doc.add_heading(f'{request.project_type} Project Report', level=1)

#     style = doc.styles['Normal']
#     font = style.font
#     font.name = request.font

#     doc.add_paragraph(generated_text)

#     file_stream = io.BytesIO()
#     doc.save(file_stream)
#     file_stream.seek(0)

#     return StreamingResponse(file_stream,
#                              media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#                              headers={
#                                  "Content-Disposition": f"attachment; filename={request.project_title.replace(' ', '_')}_project_report.docx"
#                              })



model = genai.GenerativeModel('gemini-1.5-flash')


class ProjectIdea(BaseModel):
    idea: str

@router.post("/generate_abstract/")
async def generate_abstract(idea: str = Form(...)):
    try:
        # Use OpenAI to generate the abstract
        prompt = f"Generate a 150-word abstract for the following project idea:\n\n{idea}"
        chat = model.start_chat(history=[])
        response = chat.send_message(prompt)
        abstract_text = response.text.strip()

        # Create a Word document with the abstract
        document = Document()
        document.add_heading('Project Abstract', level=1)
        document.add_paragraph(abstract_text)
        file_path = "abstract.docx"
        document.save(file_path)

        return {"abstract": abstract_text, "file_path": file_path}

    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))

@router.get("/download_abstract/")
async def download_abstract():
    file_path = "abstract.docx"
    if os.path.exists(file_path):
        return FileResponse(file_path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename="abstract.docx")
    else:
        raise HTTPException(status_code=404, detail="File not found")
